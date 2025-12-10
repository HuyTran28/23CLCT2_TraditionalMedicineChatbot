# Contains python-docx & tagging logic
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class WordExporter:
    def __init__(self, base_spacing=1.0, font_size=12, max_image_width=6.0):
        self.base_spacing = base_spacing
        self.font_size = font_size
        self.max_image_width = max_image_width  # Maximum width in inches for embedded images

    @staticmethod
    def center_y(box):
        return sum(p[1] for p in box) / len(box)
    
    def add_image_to_document(self, doc, image_data):
        """
        Add an image to the Word document with proper formatting
        
        Args:
            doc: Document object
            image_data: Dictionary with 'file_path', 'image_id', 'width', 'height'
        
        Returns:
            bool: True if image was added successfully, False otherwise
        """
        try:
            image_path = Path(image_data.get('file_path', ''))
            
            # Skip if file doesn't exist (filtered out during extraction)
            if not image_path.exists():
                logger.debug(f"Skipping image {image_data.get('image_id', 'unknown')}: file not found (likely filtered)")
                return False
            
            # Add image with caption
            image_id = image_data.get('image_id', 'unknown')
            
            # Add spacing before image
            spacing_p = doc.add_paragraph()
            spacing_p.paragraph_format.space_before = Pt(6)
            spacing_p.paragraph_format.space_after = Pt(3)
            
            # Add paragraph for image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Calculate image dimensions (maintain aspect ratio)
            original_width = image_data.get('width', 800)
            original_height = image_data.get('height', 600)

            # Validate dimensions and set defaults if invalid
            if original_width <= 0:
                logger.warning(f"Invalid width for image {image_id}, using default")
                original_width = 800
            if original_height <= 0:
                logger.warning(f"Invalid height for image {image_id}, using default")
                original_height = 600

            aspect_ratio = original_height / original_width
            
            # Adjust max width based on aspect ratio (taller images get less width)
            if aspect_ratio > 1.5:  # Very tall image
                max_width = min(4.0, self.max_image_width)
            elif aspect_ratio > 1.2:  # Tall image  
                max_width = min(5.0, self.max_image_width)
            else:  # Normal or wide image
                max_width = self.max_image_width
            
            # Calculate final dimensions
            width_inches = min(max_width, original_width / 100)
            height_inches = width_inches * aspect_ratio
            
            # Limit height to avoid overly tall images
            max_height = 8.0  # Maximum 8 inches tall
            if height_inches > max_height:
                height_inches = max_height
                width_inches = height_inches / aspect_ratio

            # Add the image
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(width_inches))
            
            # Add caption below image with better formatting
            caption_p = doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_before = Pt(3)
            caption_p.paragraph_format.space_after = Pt(6)
            caption_run = caption_p.add_run(f"[{image_id}]")
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(100, 100, 100)
            
            logger.debug(f"Added image {image_id} to document ({width_inches:.1f}x{height_inches:.1f} inches)")
            return True
            
        except FileNotFoundError:
            logger.debug(f"Image file not found: {image_data.get('image_id', 'unknown')}")
            return False
        except Exception as e:
            logger.warning(f"Failed to add image {image_data.get('image_id', 'unknown')}: {e}")
            return False

    def add_table_to_document(self, doc, table_data):
        """
        Add a table to the Word document with improved formatting
        
        Args:
            doc: Document object
            table_data: Dictionary with table structure and cell data
        
        Returns:
            bool: True if table was added successfully
        """
        try:
            table_id = table_data.get('table_id', 'unknown')
            structure = table_data.get('structure', {})
            cells = structure.get('cells', [])
            
            if not cells:
                # If no structure, add as image
                return self.add_image_to_document(doc, table_data)
            
            rows = len(cells)
            cols = len(cells[0]) if cells else 0
            
            if rows == 0 or cols == 0:
                return False
            
            # Add spacing before table
            spacing_p = doc.add_paragraph()
            spacing_p.paragraph_format.space_before = Pt(6)
            spacing_p.paragraph_format.space_after = Pt(3)
            
            # Add table title
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"[{table_id}]")
            title_run.font.bold = True
            title_run.font.size = Pt(10)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(3)
            
            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Apply table formatting
            table.autofit = False
            table.allow_autofit = False
            
            # Fill cells
            for i, row_data in enumerate(cells):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_data.get('text', '')
                    
                    # Format cell text
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(2)
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            
                    # Add cell padding
                    cell.vertical_alignment = 1  # Center vertically
            
            # Add spacing after table
            after_spacing_p = doc.add_paragraph()
            after_spacing_p.paragraph_format.space_before = Pt(3)
            after_spacing_p.paragraph_format.space_after = Pt(6)
            
            logger.debug(f"Added table {table_id} ({rows}x{cols}) to document")
            return True
            
        except Exception as e:
            logger.warning(f"Failed to add table {table_data.get('table_id', 'unknown')}: {e}")
            # Fallback: try to add as image
            try:
                return self.add_image_to_document(doc, table_data)
            except:
                return False

    @staticmethod
    def inject_break_tag(text: str) -> str:
        """Insert </break> tag after level 2 headings (e.g., 2.1, 2.2, 3.1, etc.)"""
        # Pattern matches level 2 headings: X.Y where X and Y are digits
        pattern = r"^\s*\d+\.\d+\s"
        if re.match(pattern, text.strip()):
            # Check if it's exactly level 2 (one dot, not 2.1.1)
            heading_part = text.strip().split()[0] if text.strip().split() else ""
            if heading_part.count('.') == 1:
                return text.rstrip() + " </break>"
        return text

    def write_to_word(self, data, output_path="output.docx", images=None, tables=None):
        """
        Export OCR results to Word document
        
        Args:
            data: Can be:
                - str: path to JSON file
                - list: OCR results (assumes single page)
                - dict with 'results' key: OCR results
                - dict with 'pages' key: Multi-page OCR results
            output_path: str, path to save docx
            images: Optional list of image dictionaries to embed
            tables: Optional list of table dictionaries to embed
        
        Returns:
            Path to saved document
        """
        if isinstance(data, str):
            # Assume it's a path to JSON file
            with open(data, "r", encoding="utf-8") as f:
                data = json.load(f)

        if isinstance(data, dict) and "results" in data:
            items = data["results"]
        elif isinstance(data, dict) and "pages" in data:
            # Multi-page format: {"pages": [{page_num: 1, results: [...]}, ...]}
            return self._write_multipage_to_word(data["pages"], output_path, images, tables)
        elif isinstance(data, list):
            items = data
        elif isinstance(data, dict):
            # For layout export: expects keys 'text', 'images', 'tables'
            doc = Document()
            if data.get("text"):
                doc.add_paragraph(data["text"])
            # Images and tables can be handled here if needed
            doc.save(output_path)
            print("Saved:", output_path)
            return output_path
        else:
            raise ValueError("Unsupported data format for export")

        # OCR results export - handle both with and without page_id
        pages = {}
        for it in items:
            page_id = it.get("page_id", 1)
            pages.setdefault(page_id, []).append(it)

        doc = Document()

        for page_index, (page_id, page_items) in enumerate(pages.items()):
            # Separate text items, image items, and table items
            text_items = [it for it in page_items if it.get("element_type") not in ["image", "table"]]
            image_items = [it for it in page_items if it.get("element_type") == "image"]
            table_items = [it for it in page_items if it.get("element_type") == "table"]
            
            # Sort text items by position
            text_items = sorted(
                text_items,
                key=lambda t: (self.center_y(t.get("box", [[0,0]])), t.get("box", [[0,0]])[0][0])
            )

            last_y = None
            last_height = 12

            for it in text_items:
                # Skip headers and footers if marked
                if it.get("skip", False):
                    continue
                
                text = it.get("text", "")
                box = it.get("box", [[0, 0], [0, 0], [0, 0], [0, 0]])
                element_type = it.get("element_type", "paragraph")

                # Apply break tag for level 2 headings
                text = self.inject_break_tag(text)

                y = self.center_y(box)
                height = abs(box[2][1] - box[0][1])

                if last_y is None:
                    spacing = self.base_spacing
                else:
                    gap = y - last_y
                    spacing = max(self.base_spacing, gap / last_height)

                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(self.font_size)
                
                # Apply heading formatting
                if element_type == "heading":
                    run.font.bold = True
                    heading_level = it.get("heading_level", 0)
                    if heading_level == 1:
                        run.font.size = Pt(16)
                    elif heading_level == 2:
                        run.font.size = Pt(14)
                    elif heading_level >= 3:
                        run.font.size = Pt(13)

                p.paragraph_format.line_spacing = spacing

                last_y = y
                last_height = max(height, 8)
            
            # Add images for this page
            if image_items:
                added_count = 0
                for img_item in image_items:
                    if self.add_image_to_document(doc, img_item):
                        added_count += 1
                if added_count > 0:
                    logger.info(f"Added {added_count}/{len(image_items)} images for page {page_id}")
            
            # Add tables for this page
            if table_items:
                added_count = 0
                for tbl_item in table_items:
                    if self.add_table_to_document(doc, tbl_item):
                        added_count += 1
                if added_count > 0:
                    logger.info(f"Added {added_count}/{len(table_items)} tables for page {page_id}")

            if page_index < len(pages) - 1:
                doc.add_page_break()

        doc.save(output_path)
        print("Saved:", output_path)
        return output_path
    
    def _write_multipage_to_word(self, pages_data, output_path="output.docx", images=None, tables=None):
        """
        Export multi-page OCR results to Word document
        
        Args:
            pages_data: List of dicts with 'page_num' and 'results' keys
            output_path: str, path to save docx
            images: Optional list of image dictionaries to embed
            tables: Optional list of table dictionaries to embed
        
        Returns:
            Path to saved document
        """
        doc = Document()
        
        # Sort pages by page number
        pages_data = sorted(pages_data, key=lambda p: p.get("page_num", 1))
        
        # Group images and tables by page
        images_by_page = {}
        tables_by_page = {}
        
        if images:
            for img in images:
                page_num = img.get("page_num", 1)
                images_by_page.setdefault(page_num, []).append(img)
        
        if tables:
            for tbl in tables:
                page_num = tbl.get("page_num", 1)
                tables_by_page.setdefault(page_num, []).append(tbl)
        
        for page_idx, page_data in enumerate(pages_data):
            page_num = page_data.get("page_num", page_idx + 1)
            results = page_data.get("results", [])
            
            if not results:
                continue
            
            # Separate text, image, and table elements
            text_results = [r for r in results if r.get("element_type") not in ["image", "table"]]
            image_results = [r for r in results if r.get("element_type") == "image"]
            table_results = [r for r in results if r.get("element_type") == "table"]
            
            # Sort text results by vertical position, then horizontal
            text_results = sorted(
                text_results,
                key=lambda t: (self.center_y(t.get("box", [[0,0]])), t.get("box", [[0,0]])[0][0])
            )
            
            last_y = None
            last_height = 12
            
            for it in text_results:
                # Skip headers and footers
                if it.get("skip", False):
                    continue
                
                text = it.get("text", "")
                box = it.get("box", [[0, 0], [0, 0], [0, 0], [0, 0]])
                element_type = it.get("element_type", "paragraph")
                
                # Apply break tag
                text = self.inject_break_tag(text)
                
                y = self.center_y(box)
                height = abs(box[2][1] - box[0][1])
                
                if last_y is None:
                    spacing = self.base_spacing
                else:
                    gap = y - last_y
                    spacing = max(self.base_spacing, gap / last_height)
                
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(self.font_size)
                
                # Apply heading formatting
                if element_type == "heading":
                    run.font.bold = True
                    heading_level = it.get("heading_level", 0)
                    if heading_level == 1:
                        run.font.size = Pt(16)
                    elif heading_level == 2:
                        run.font.size = Pt(14)
                    elif heading_level >= 3:
                        run.font.size = Pt(13)
                
                p.paragraph_format.line_spacing = spacing
                
                last_y = y
                last_height = max(height, 8)
            
            # Add images from results
            added_from_results = 0
            for img_item in image_results:
                if self.add_image_to_document(doc, img_item):
                    added_from_results += 1
            
            # Add images from separate image list
            added_from_list = 0
            if page_num in images_by_page:
                for img in images_by_page[page_num]:
                    if self.add_image_to_document(doc, img):
                        added_from_list += 1
            
            # Add tables from results
            added_tables_from_results = 0
            for tbl_item in table_results:
                if self.add_table_to_document(doc, tbl_item):
                    added_tables_from_results += 1
            
            # Add tables from separate table list
            added_tables_from_list = 0
            if page_num in tables_by_page:
                for tbl in tables_by_page[page_num]:
                    if self.add_table_to_document(doc, tbl):
                        added_tables_from_list += 1
            
            # Log summary
            total_images = added_from_results + added_from_list
            total_tables = added_tables_from_results + added_tables_from_list
            
            if total_images > 0:
                logger.info(f"Page {page_num}: Added {total_images} images")
            if total_tables > 0:
                logger.info(f"Page {page_num}: Added {total_tables} tables")
            
            # Add page break between pages (except last page)
            if page_idx < len(pages_data) - 1:
                doc.add_page_break()
        
        doc.save(output_path)
        print(f"Saved {len(pages_data)} pages to: {output_path}")
        return output_path

    

    def markdown_to_word(self, markdown_text, output_path="output.docx", images=None):
        """
        Convert markdown text to Word document with native tables and formatting.
        """
        doc = Document()
        
        # --- Helper: Image Lookup ---
        image_map = {img.get('image_id', ''): img for img in images} if images else {}

        # --- Helper: Apply formatting (Bold/Italic) ---
        def add_formatted_text(paragraph, text):
            # Split text by bold (**...**) and italic (*...*) markers
            # The regex keeps the delimiters in the list so we can identify them
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            
            for part in parts:
                if not part:
                    continue
                
                run = paragraph.add_run()
                if part.startswith('**') and part.endswith('**'):
                    run.text = part[2:-2]
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run.text = part[1:-1]
                    run.font.italic = True
                else:
                    run.text = part
                    
        # --- Helper: Process Image Placeholder ---
        def process_image_placeholder(line):
            match = re.search(r'\[IMAGE_PLACEHOLDER_(\d+)\]', line)
            if match:
                img_num = match.group(1)
                img_id = f'img_{img_num}'
                if img_id in image_map:
                    self.add_image_to_document(doc, image_map[img_id])
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[IMAGE {img_num}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
            return False

        # --- Main Parsing Loop ---
        lines = markdown_text.split('\n')
        table_buffer = [] # To store table rows temporarily
        in_table = False
        
        for i, line in enumerate(lines):
            clean_line = line.strip()
            
            # 1. Handle Table Logic
            if clean_line.startswith('|'):
                # It's a table row
                if not in_table:
                    # Check if the NEXT line is a separator (|---|) to confirm it's a header
                    if i + 1 < len(lines) and '---' in lines[i+1]:
                        in_table = True
                        table_buffer = [line] # Start buffer
                        continue # Skip processing this line as normal text
                else:
                    # If we are already in a table, just add the line
                    # Ignore the separator line (e.g. |---|---|)
                    if '---' in line:
                        continue 
                    table_buffer.append(line)
                    continue
            
            # If we were in a table, but this line is NOT a table line (or empty)
            if in_table:
                # Render the buffered table now
                if table_buffer:
                    # Calculate rows and columns
                    rows_data = [[cell.strip() for cell in row.strip('|').split('|')] for row in table_buffer]
                    num_rows = len(rows_data)
                    num_cols = len(rows_data[0]) if rows_data else 0
                    
                    if num_rows > 0 and num_cols > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'
                        
                        for r_idx, row_data in enumerate(rows_data):
                            row_cells = table.rows[r_idx].cells
                            for c_idx, cell_text in enumerate(row_data):
                                if c_idx < len(row_cells):
                                    # Convert HTML <br> tags to proper line breaks
                                    cell_text = cell_text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
                                    
                                    # Split by newlines and add as separate paragraphs for proper line breaks
                                    lines_in_cell = cell_text.split('\n')
                                    p = row_cells[c_idx].paragraphs[0]
                                    add_formatted_text(p, lines_in_cell[0])
                                    
                                    # Add additional lines as new paragraphs
                                    for cell_line in lines_in_cell[1:]:
                                        new_p = row_cells[c_idx].add_paragraph()
                                        add_formatted_text(new_p, cell_line)
                
                # Reset table state
                in_table = False
                table_buffer = []
                
                # If this current line was empty, we are done with the table. 
                if not clean_line:
                    continue

            # 2. Handle Images
            if '[IMAGE_PLACEHOLDER_' in line:
                if process_image_placeholder(line):
                    continue

            # 3. Skip empty lines (outside of tables)
            if not clean_line:
                continue

            # 4. Handle Headings
            if line.startswith('#'):
                level = len(line.split(' ')[0]) # Count hashes
                text = line[level:].strip()
                if 1 <= level <= 4:
                    p = doc.add_paragraph()
                    add_formatted_text(p, text) # Support bold in headers
                    # Basic Styling based on level
                    run = p.runs[0] # Apply size to first run or loop all
                    p.style = f'Heading {level}'
                continue

            # 5. Handle Lists
            if line.strip().startswith(('- ', '* ')):
                p = doc.add_paragraph(style='List Bullet')
                text = line.strip()[2:]
                add_formatted_text(p, text)
                p.paragraph_format.left_indent = Pt(20)
                continue
                
            elif re.match(r'^\s*\d+\.\s', line):
                # Numbered list
                text = re.sub(r'^\s*\d+\.\s', '', line)
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, text)
                p.paragraph_format.left_indent = Pt(20)
                continue

            # 6. Handle break tags (skip them as they're document structure markers)
            if clean_line == '</break>':
                # Page/section break handling - could add a page break here if needed
                continue
            
            # 7. Handle HTML <br> tags in regular paragraphs
            if '<br>' in line or '<br/>' in line or '<br />' in line:
                line = line.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
                # Split by newlines and add as paragraph with internal line breaks
                p = doc.add_paragraph()
                lines_in_para = line.split('\n')
                for idx, para_line in enumerate(lines_in_para):
                    if idx > 0:
                        p.add_run('\n')  # Line break within paragraph
                    add_formatted_text(p, para_line)
            else:
                # Regular Paragraph
                p = doc.add_paragraph()
                add_formatted_text(p, line)

        doc.save(output_path)
        return output_path

if __name__ == "__main__":
    # Example usage
    exporter = WordExporter()
    exporter.write_to_word("vietnamese_ocr_results.json", "output.docx")